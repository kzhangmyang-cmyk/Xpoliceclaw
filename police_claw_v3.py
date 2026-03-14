#!/usr/bin/env python3
# ═══════════════════════════════════════════════════════════════
#  Police Claw v3.0 — Enterprise Security Scanner
#  AI Agent / Variant Lobster / Skills 全域安全隐私审计
#  Single-file Architecture · Immutable Logic · Fixed Output
# ═══════════════════════════════════════════════════════════════

"""
Police Claw v3.0

核心使命：
  对套客（Socket）/ 变体龙虾（Variant Lobster）/ 智能体 / Skills
  执行全权限安全隐私检查。

升级要点（v2 → v3）：
  ① 检查项从 16 项扩展到 42 项
  ② 检查项按 7 大安全域分类
  ③ 新增网络流量监控层（识别数据是否上传云端）
  ④ 新增文件系统实时监控层（检测 wallet / ssh / cookie 读取）
  ⑤ 新增模型 API 调用检测层（识别数据是否进入模型上下文）
  ⑥ Risk Engine 升级为加权评分 + 置信度
  ⑦ Report Writer 输出带分类的结构化报告

架构：
  Collector       → 采集进程/命令行/文件句柄/网络连接/环境变量/DNS
  Traffic Monitor → 网络流量特征分析
  FS Monitor      → 文件系统敏感路径访问检测
  Model Monitor   → 模型 API 调用与上下文泄露检测
  Signal Engine   → 42 项信号识别
  Risk Engine     → 加权评分 + 分类 + 置信度
  Report Writer   → JSON + DOCX 分类报告

运行：
  pip install psutil python-docx
  python police_claw_v3.py [output_dir]
"""

import json, os, re, sys, platform, datetime, hashlib, socket, subprocess, ipaddress
from pathlib import Path
from collections import defaultdict
from typing import Any

try:
    import psutil
except ImportError:
    sys.exit("[!] 需要 psutil: pip install psutil --break-system-packages")

try:
    from docx import Document as DocxDocument
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def _prepare_stdio():
    for stream in (sys.stdout, sys.stderr):
        if hasattr(stream, "reconfigure"):
            try:
                stream.reconfigure(encoding="utf-8", errors="replace")
            except ValueError:
                continue


_prepare_stdio()

# ━━━━━━━━━━━━━━━━━━ 版本 & ID ━━━━━━━━━━━━━━━━━━

VERSION = "3.0.0"
SCAN_TS = datetime.datetime.now()
SCAN_ID = hashlib.sha256(
    f"{platform.node()}-{SCAN_TS.isoformat()}".encode()
).hexdigest()[:12].upper()

# ━━━━━━━━━━━━━━━━━━ 安全域定义 ━━━━━━━━━━━━━━━━━━
# 7 大安全域，每个域包含多项检查

DOMAINS = {
    "credential": {
        "name": "凭证与身份安全",
        "icon": "🔐",
        "weight": 1.5,  # 高权重
    },
    "transaction": {
        "name": "交易与金融安全",
        "icon": "💰",
        "weight": 1.4,
    },
    "behavior": {
        "name": "用户行为追踪",
        "icon": "👁️",
        "weight": 1.2,
    },
    "system": {
        "name": "系统权限与控制",
        "icon": "⚙️",
        "weight": 1.3,
    },
    "data": {
        "name": "数据采集与外泄",
        "icon": "📡",
        "weight": 1.4,
    },
    "model": {
        "name": "模型与 AI 上下文",
        "icon": "🧠",
        "weight": 1.2,
    },
    "audit": {
        "name": "审计与合规",
        "icon": "🛡️",
        "weight": 1.0,
    },
}

# ━━━━━━━━━━━━━━━━━━ 42 项检查定义 ━━━━━━━━━━━━━━━━━━

CHECK_ITEMS = [
    # ── 凭证与身份安全 (7项) ──
    ("cred_password",       "credential", "抓取账户密码",              "检测进程是否访问密码存储 / keychain / 浏览器密码数据库"),
    ("cred_ssh_keys",       "credential", "读取 SSH 密钥",            "检测对 ~/.ssh/id_rsa, id_ed25519 等私钥文件的访问"),
    ("cred_api_tokens",     "credential", "窃取 API Token",           "检测环境变量和配置文件中 API 密钥的暴露与读取"),
    ("cred_cookies",        "credential", "抓取浏览器 Cookie",        "检测对 Chrome/Firefox/Safari cookie 数据库的访问"),
    ("cred_wallet",         "credential", "读取加密钱包",             "检测对 wallet.dat / MetaMask vault / 助记词文件的访问"),
    ("cred_2fa",            "credential", "窃取 2FA 凭证",            "检测对 TOTP 种子、Authenticator 数据库的访问"),
    ("cred_cert",           "credential", "读取证书与私钥",           "检测对 .pem / .p12 / .pfx / TLS 私钥文件的访问"),

    # ── 交易与金融安全 (4项) ──
    ("txn_unauthorized",    "transaction", "未授权进行交易",           "监控未授权的金融交易进程或到交易所的网络连接"),
    ("txn_crypto",          "transaction", "加密货币自动交易",         "检测与 DEX / CEX API 的未授权交互"),
    ("txn_payment",         "transaction", "篡改支付信息",             "检测对支付网关请求的劫持或中间人行为"),
    ("txn_mining",          "transaction", "挖矿行为",                 "检测 xmrig / cpuminer 等挖矿进程及矿池连接"),

    # ── 用户行为追踪 (7项) ──
    ("beh_search",          "behavior", "抓取搜索行为",               "检测浏览器历史 / 搜索日志 / 自动补全数据的异常读取"),
    ("beh_code",            "behavior", "追踪代码编写",               "监控 .git / IDE 工作区 / 代码仓库的异常扫描"),
    ("beh_debug",           "behavior", "追踪调试行为",               "检测 gdb / lldb / strace / pdb 等调试工具的异常调用"),
    ("beh_keylog",          "behavior", "键盘记录",                    "检测键盘输入捕获进程 / 输入法监控 / IME 注入"),
    ("beh_screen",          "behavior", "屏幕截图与录制",              "检测屏幕捕获 / 截图进程 / 远程桌面未授权共享"),
    ("beh_clipboard",       "behavior", "剪贴板监控",                  "检测剪贴板内容持续读取或劫持（尤其是加密地址替换）"),
    ("beh_operation_log",   "behavior", "操作记录抓取",                "检测用户操作轨迹的异常采集与外传"),

    # ── 系统权限与控制 (6项) ──
    ("sys_root",            "system", "root/SYSTEM 权限运行",          "检查 AI 进程是否以最高权限运行"),
    ("sys_persistence",     "system", "持久化驻留",                    "检测 crontab / 启动项 / systemd service 的异常注册"),
    ("sys_process_inject",  "system", "进程注入",                      "检测 DLL 注入 / ptrace attach / LD_PRELOAD 劫持"),
    ("sys_driver",          "system", "内核模块/驱动加载",              "检测未授权的内核模块或驱动程序加载"),
    ("sys_firewall",        "system", "防火墙规则篡改",                "检测 iptables / Windows Firewall 规则的异常修改"),
    ("sys_dns",             "system", "DNS 劫持",                      "检测 DNS 配置篡改 / hosts 文件修改 / 异常 DNS 服务器"),

    # ── 数据采集与外泄 (9项) ──
    ("data_file_read",      "data", "读取私人文件",                    "检测对 Documents / Photos / Downloads 的异常批量访问"),
    ("data_file_content",   "data", "抓取文件内容",                    "检测文件索引 / OCR / 文本提取 / PDF 解析进程"),
    ("data_cloud_upload",   "data", "上传云端文件",                    "监控云同步进程和 rclone / s3 cp / gsutil 等上传命令"),
    ("data_idle_exfil",     "data", "待机偷跑数据",                    "检测休眠进程的隐蔽网络外传行为"),
    ("data_stream",         "data", "平台可见全部数据流",               "监控大规模外传连接 / 遥测 / 分析进程"),
    ("data_dns_tunnel",     "data", "DNS 隧道外泄",                    "检测通过 DNS TXT/CNAME 记录进行的数据编码外传"),
    ("data_steganography",  "data", "隐写术数据外泄",                  "检测图片 / 音频文件中嵌入隐藏数据的外传行为"),
    ("data_usb",            "data", "USB/外接设备数据拷贝",             "检测向 USB 设备或外接存储的异常大规模数据传输"),
    ("data_backup_exfil",   "data", "备份文件外泄",                    "检测对系统备份 / Time Machine / 快照文件的异常访问"),

    # ── 模型与 AI 上下文 (5项) ──
    ("model_context",       "model", "数据进入模型上下文",              "检测 LLM / RAG 管道是否将敏感数据纳入推理上下文"),
    ("model_prompt",        "model", "Prompt 抓取",                    "检测 Prompt 缓存 / system prompt / 对话历史的异常读取"),
    ("model_finetune",      "model", "用户数据用于微调",                "检测本地数据是否被用于模型 fine-tune / LoRA 训练"),
    ("model_embedding",     "model", "敏感数据向量化",                  "检测私人文档被 embedding 化存入向量数据库"),
    ("model_api_leak",      "model", "模型 API 调用泄露数据",           "检测向外部模型 API 发送的请求中是否包含敏感信息"),

    # ── 审计与合规 (4项) ──
    ("audit_system",        "audit", "安全审计体系",                    "验证 auditd / SIEM / 安全日志系统是否正常运行"),
    ("audit_log_tamper",    "audit", "审计日志篡改",                    "检测安全日志的异常删除 / 截断 / 权限变更"),
    ("audit_compliance",    "audit", "合规性缺失",                      "检查 GDPR / CCPA / 个人信息保护法等合规措施是否到位"),
    ("audit_leak_risk",     "audit", "综合数据泄露风险",                "多类目同时触发时的复合风险评估"),
]

# ━━━━━━━━━━━━━━━━━━ 检测关键词库 ━━━━━━━━━━━━━━━━━━

KW = {
    "passwords":    ["passwd", "shadow", "keychain", "credential", "login_data",
                     "password", "logins.json", "web data", "login keychain"],
    "ssh":          [".ssh", "id_rsa", "id_ed25519", "id_ecdsa", "authorized_keys",
                     "known_hosts", "ssh_config"],
    "api_tokens":   ["api_key", "secret_key", "access_token", "bearer", ".env",
                     ".npmrc", ".pypirc", "credentials.json", "service_account"],
    "cookies":      ["cookies.sqlite", "cookies", "cookie", "session_store",
                     "local storage", "indexed_db"],
    "wallet":       ["wallet.dat", "metamask", "vault", "mnemonic", "seed_phrase",
                     "keystore", "electrum", "bitcoin", "ethereum"],
    "2fa":          ["authenticator", "totp", "2fa", "otp_secret", "aegis",
                     "andotp", "google_authenticator"],
    "certs":        [".pem", ".p12", ".pfx", ".key", ".crt", "private_key",
                     "tls_cert", "ssl_cert", "ca-bundle"],
    "trade":        ["trade", "broker", "exchange", "coinbase", "binance",
                     "robinhood", "etrade", "kraken", "ftx", "dex"],
    "mining":       ["xmrig", "cpuminer", "cgminer", "bfgminer", "ethminer",
                     "phoenix", "nbminer", "stratum+tcp", "stratum+ssl",
                     "mining_pool", "nicehash"],
    "search":       ["history.db", "places.sqlite", "browsing_history",
                     "search_history", "autocomplete", "omnibox"],
    "code":         [".git", "vscode", "workspace", "repo", ".idea",
                     "jetbrains", "sublime", "source_code"],
    "debug":        ["gdb", "lldb", "strace", "dtrace", "pdb", "debug_log",
                     "breakpoint", "valgrind", "perf"],
    "keylog":       ["keylogger", "keylog", "input_capture", "keyboard_hook",
                     "input_monitor", "ime_inject"],
    "screen":       ["screenshot", "screen_capture", "screen_record",
                     "vnc", "rdp", "teamviewer", "anydesk", "screenlog"],
    "clipboard":    ["clipboard", "pbpaste", "xclip", "xsel",
                     "clipboard_monitor", "clip_hijack"],
    "cloud":        ["dropbox", "onedrive", "googledrive", "icloud",
                     "syncthing", "rclone", "mega", "nextcloud", "box"],
    "upload_cmd":   ["upload", "s3 cp", "gsutil cp", "rclone copy",
                     "az storage", "aws s3", "scp ", "rsync"],
    "ai_proc":      ["ollama", "llama", "gpt", "claude", "langchain",
                     "autogpt", "babyagi", "copilot", "openai",
                     "anthropic", "huggingface", "transformers"],
    "model_api":    ["api.openai.com", "api.anthropic.com", "api.cohere",
                     "huggingface.co/api", "api.mistral", "generativelanguage"],
    "embedding":    ["embedding", "vector_store", "chroma", "pinecone",
                     "weaviate", "qdrant", "faiss", "milvus"],
    "finetune":     ["fine_tune", "finetune", "lora", "qlora", "peft",
                     "training_data", "dataset_prepare"],
    "prompt":       ["prompt_log", "prompt_cache", "system_prompt",
                     "conversation_history", "chat_history", "instruction"],
    "audit":        ["auditd", "audit.log", "syslog", "security.log",
                     "ossec", "wazuh", "falco", "siem", "splunk"],
    "persist":      ["crontab", "cron.d", "systemd", "launchd", "rc.local",
                     "startup", "autorun", "registry\\run", "init.d"],
    "inject":       ["ptrace", "ld_preload", "dll_inject", "code_inject",
                     "process_hollow", "reflective_load"],
    "dns_tunnel":   ["iodine", "dns2tcp", "dnscat", "dns_tunnel",
                     "dns_exfil", "covert_dns"],
    "telemetry":    ["telemetry", "analytics", "beacon", "tracker",
                     "pixel", "heartbeat", "phone_home"],
    "private":      ["documents", "desktop", "downloads", "pictures",
                     "photos", "personal", "diary", "medical", "tax"],
    "backup":       ["time_machine", "backup", "snapshot", ".bak",
                     "recovery", "system_restore"],
}

# ━━━━━━━━━━━━━━━━━━ Collector ━━━━━━━━━━━━━━━━━━

class Collector:
    """采集系统运行时信息"""

    def __init__(self):
        self.processes = []
        self.connections = []
        self.open_files = []
        self.env_signals = []
        self.dns_servers = []
        self.listening_ports = []

    def collect_all(self):
        self._collect_processes()
        self._collect_connections()
        self._collect_env()
        self._collect_dns()
        return self

    def _collect_processes(self):
        for proc in psutil.process_iter(
            ["pid", "name", "cmdline", "username", "status",
             "open_files", "create_time"]
        ):
            try:
                info = proc.info
                entry = {
                    "pid": info["pid"],
                    "name": (info["name"] or "").lower(),
                    "cmdline": " ".join(info["cmdline"] or []).lower(),
                    "username": info["username"] or "",
                    "status": info["status"] or "",
                    "open_files": [],
                    "age_hours": 0,
                }
                if info.get("create_time"):
                    age = (SCAN_TS.timestamp() - info["create_time"]) / 3600
                    entry["age_hours"] = round(age, 1)
                if info["open_files"]:
                    entry["open_files"] = [f.path for f in info["open_files"]]
                    self.open_files.extend(entry["open_files"])
                self.processes.append(entry)
            except (psutil.NoSuchProcess, psutil.AccessDenied,
                    psutil.ZombieProcess):
                continue

    def _collect_connections(self):
        proc_map = {p["pid"]: p for p in self.processes}
        try:
            for conn in psutil.net_connections(kind="inet"):
                proc = proc_map.get(conn.pid or -1, {})
                entry = {
                    "pid": conn.pid,
                    "local": f"{conn.laddr.ip}:{conn.laddr.port}" if conn.laddr else "",
                    "remote": f"{conn.raddr.ip}:{conn.raddr.port}" if conn.raddr else "",
                    "status": conn.status,
                    "family": str(conn.family),
                    "process_name": proc.get("name", ""),
                    "process_cmdline": proc.get("cmdline", ""),
                }
                self.connections.append(entry)
                if conn.status == "LISTEN":
                    self.listening_ports.append(conn.laddr.port if conn.laddr else 0)
        except (psutil.AccessDenied, OSError):
            pass

    def _collect_env(self):
        suspect = ["API_KEY", "SECRET", "TOKEN", "PASSWORD", "CREDENTIAL",
                    "AWS_", "AZURE_", "GCP_", "OPENAI", "ANTHROPIC",
                    "DATABASE_URL", "REDIS_URL", "MONGO", "PRIVATE_KEY",
                    "WALLET", "MNEMONIC", "SEED"]
        for key in os.environ:
            for pat in suspect:
                if pat in key.upper():
                    self.env_signals.append({
                        "key": key,
                        "hint": f"环境变量含敏感关键词: {pat}",
                    })
                    break

    def _collect_dns(self):
        seen = set()

        def add_server(value: str):
            value = value.strip()
            if not value or value in seen:
                return
            seen.add(value)
            self.dns_servers.append(value)

        resolv = Path("/etc/resolv.conf")
        if resolv.exists():
            try:
                for line in resolv.read_text().splitlines():
                    if line.strip().startswith("nameserver"):
                        add_server(line.strip().split()[-1])
            except Exception:
                pass

        if platform.system() == "Windows":
            commands = [
                [
                    "powershell",
                    "-NoProfile",
                    "-Command",
                    "Get-DnsClientServerAddress -AddressFamily IPv4 | Select-Object -ExpandProperty ServerAddresses",
                ],
                ["ipconfig", "/all"],
            ]
            for command in commands:
                try:
                    output = subprocess.check_output(
                        command,
                        text=True,
                        encoding="utf-8",
                        errors="ignore",
                        timeout=6,
                    )
                except Exception:
                    continue

                for match in re.findall(r"(?:\d{1,3}\.){3}\d{1,3}", output):
                    add_server(match)

                if self.dns_servers:
                    break


# ━━━━━━━━━━━━━━━━━━ Traffic Monitor ━━━━━━━━━━━━━━━━━━

class TrafficMonitor:
    """网络流量特征分析"""

    def __init__(self, collector: Collector):
        self.c = collector

    def analyze(self) -> dict:
        results = {
            "outbound_count": 0,
            "cloud_endpoints": [],
            "model_api_endpoints": [],
            "suspicious_ports": [],
            "dns_anomaly": False,
        }
        established = [c for c in self.c.connections
                       if c["status"] == "ESTABLISHED" and c["remote"]]
        results["outbound_count"] = len(established)

        for conn in established:
            remote = conn["remote"]
            conn_text = " ".join(
                [remote, conn.get("process_name", ""), conn.get("process_cmdline", "")]
            ).lower()
            # 云端端点
            for kw in KW["cloud"] + KW["upload_cmd"]:
                if kw in conn_text:
                    endpoint = remote
                    if conn.get("process_name"):
                        endpoint = f"{remote} [{conn['process_name']}]"
                    results["cloud_endpoints"].append(endpoint)
            # 模型 API 端点
            for kw in KW["model_api"]:
                if kw in conn_text:
                    endpoint = remote
                    if conn.get("process_name"):
                        endpoint = f"{remote} [{conn['process_name']}]"
                    results["model_api_endpoints"].append(endpoint)
            # 可疑端口 (非标准)
            try:
                port = int(remote.split(":")[-1])
                if port not in (80, 443, 22, 53, 8080, 8443, 3306, 5432):
                    results["suspicious_ports"].append(
                        {"remote": remote, "port": port, "pid": conn["pid"]})
            except (ValueError, IndexError):
                pass

        # DNS 异常检测
        known_dns = {"8.8.8.8", "8.8.4.4", "1.1.1.1", "1.0.0.1",
                      "208.67.222.222", "208.67.220.220", "127.0.0.53"}
        for dns in self.c.dns_servers:
            try:
                if ipaddress.ip_address(dns).is_private:
                    continue
            except ValueError:
                continue
            if dns not in known_dns:
                results["dns_anomaly"] = True
                break

        return results


# ━━━━━━━━━━━━━━━━━━ FS Monitor ━━━━━━━━━━━━━━━━━━

class FSMonitor:
    """文件系统敏感路径访问检测"""

    SENSITIVE_ZONES = {
        "wallet":  [".bitcoin", ".ethereum", "metamask", "wallet",
                    "electrum", "keystore"],
        "ssh":     [".ssh", "id_rsa", "id_ed25519", "known_hosts"],
        "cookie":  ["cookies.sqlite", "cookies", "cookie_db",
                    "login_data", "logins.json"],
        "cert":    [".pem", ".p12", ".pfx", ".key", "private_key"],
        "env":     [".env", ".npmrc", ".pypirc", "credentials.json"],
        "backup":  ["time_machine", ".bak", "backup", "snapshot"],
    }

    def __init__(self, collector: Collector):
        self.c = collector

    def analyze(self) -> dict:
        results = defaultdict(list)
        for fpath in self.c.open_files:
            fl = fpath.lower()
            for zone, keywords in self.SENSITIVE_ZONES.items():
                for kw in keywords:
                    if kw in fl:
                        results[zone].append(fpath)
                        break
        return dict(results)


# ━━━━━━━━━━━━━━━━━━ Model Monitor ━━━━━━━━━━━━━━━━━━

class ModelMonitor:
    """模型 API 调用与上下文泄露检测"""

    def __init__(self, collector: Collector, traffic: dict):
        self.c = collector
        self.traffic = traffic

    def analyze(self) -> dict:
        results = {
            "model_procs": [],
            "api_connections": self.traffic.get("model_api_endpoints", []),
            "embedding_procs": [],
            "finetune_procs": [],
            "prompt_files": [],
        }
        for p in self.c.processes:
            text = f"{p['name']} {p['cmdline']}"
            if any(kw in text for kw in KW["ai_proc"]):
                results["model_procs"].append(
                    {"pid": p["pid"], "name": p["name"]})
            if any(kw in text for kw in KW["embedding"]):
                results["embedding_procs"].append(
                    {"pid": p["pid"], "name": p["name"]})
            if any(kw in text for kw in KW["finetune"]):
                results["finetune_procs"].append(
                    {"pid": p["pid"], "name": p["name"]})

        for f in self.c.open_files:
            fl = f.lower()
            if any(kw in fl for kw in KW["prompt"]):
                results["prompt_files"].append(f)

        return results


# ━━━━━━━━━━━━━━━━━━ Signal Engine ━━━━━━━━━━━━━━━━━━

class SignalEngine:
    """从采集数据中提取 42 项安全信号"""

    def __init__(self, collector, traffic, fs_data, model_data):
        self.c = collector
        self.traffic = traffic
        self.fs = fs_data
        self.model = model_data

    def _match_proc(self, keywords):
        hits = []
        for p in self.c.processes:
            text = f"{p['name']} {p['cmdline']}"
            for kw in keywords:
                if kw in text:
                    hits.append({"type": "process", "pid": p["pid"],
                                 "name": p["name"], "match": kw})
                    break
        return hits

    def _match_files(self, keywords):
        return [{"type": "file", "path": f}
                for f in self.c.open_files
                if any(k in f.lower() for k in keywords)]

    def _match_env(self, keywords):
        return [{"type": "env", **e}
                for e in self.c.env_signals
                if any(k in e["key"].upper() for k in keywords)]

    def analyze(self) -> dict:
        S = {}

        # ── Credential Domain ──
        S["cred_password"]  = self._match_files(KW["passwords"]) + \
                              self._match_proc(["keylog", "mimikatz", "lazagne", "credential_dump"])
        S["cred_ssh_keys"]  = self._match_files(KW["ssh"])
        S["cred_api_tokens"]= self._match_files(KW["api_tokens"]) + \
                              self._match_env(["API_KEY", "SECRET", "TOKEN"])
        S["cred_cookies"]   = self._match_files(KW["cookies"])
        S["cred_wallet"]    = self._match_files(KW["wallet"]) + \
                              [{"type": "fs_monitor", "zone": "wallet", "paths": self.fs.get("wallet", [])}] \
                              if self.fs.get("wallet") else self._match_files(KW["wallet"])
        S["cred_2fa"]       = self._match_files(KW["2fa"])
        S["cred_cert"]      = self._match_files(KW["certs"]) + \
                              [{"type": "fs_monitor", "zone": "cert", "paths": self.fs.get("cert", [])}] \
                              if self.fs.get("cert") else self._match_files(KW["certs"])

        # ── Transaction Domain ──
        S["txn_unauthorized"]= self._match_proc(KW["trade"])
        S["txn_crypto"]      = [{"type": "network", "endpoint": ep}
                                for ep in self.traffic.get("cloud_endpoints", [])
                                if any(x in ep for x in ["binance", "coinbase", "kraken", "dex"])]
        S["txn_payment"]     = self._match_proc(["payment", "stripe", "paypal", "alipay", "wechat_pay"])
        S["txn_mining"]      = self._match_proc(KW["mining"])

        # ── Behavior Domain ──
        S["beh_search"]      = self._match_files(KW["search"])
        S["beh_code"]        = self._match_files(KW["code"])
        S["beh_debug"]       = self._match_proc(KW["debug"]) + self._match_files(KW["debug"])
        S["beh_keylog"]      = self._match_proc(KW["keylog"])
        S["beh_screen"]      = self._match_proc(KW["screen"])
        S["beh_clipboard"]   = self._match_proc(KW["clipboard"])
        S["beh_operation_log"]= self._match_proc(["activity_monitor", "input_log", "user_track"])

        # ── System Domain ──
        sys_elevated = []
        for p in self.c.processes:
            if p["username"] in ("root", "SYSTEM", "NT AUTHORITY\\SYSTEM"):
                if any(ai in p["name"] for ai in KW["ai_proc"]):
                    sys_elevated.append({"type": "elevated_ai", "pid": p["pid"],
                                         "name": p["name"], "user": p["username"]})
        S["sys_root"]        = sys_elevated
        S["sys_persistence"] = self._match_proc(KW["persist"]) + self._match_files(KW["persist"])
        S["sys_process_inject"]= self._match_proc(KW["inject"])
        S["sys_driver"]      = self._match_proc(["insmod", "modprobe", "kext", "driver_load"])
        S["sys_firewall"]    = self._match_proc(["iptables", "ufw", "netsh", "firewall", "nftables"])
        S["sys_dns"]         = [{"type": "dns_anomaly", "servers": self.c.dns_servers}] \
                               if self.traffic.get("dns_anomaly") else []

        # ── Data Domain ──
        S["data_file_read"]  = self._match_files(KW["private"])
        S["data_file_content"]= self._match_proc(["file_index", "text_extract", "ocr",
                                                    "pdf_parse", "doc_scan", "content_scan"])
        S["data_cloud_upload"]= self._match_proc(KW["cloud"] + KW["upload_cmd"]) + \
                                [{"type": "traffic", "endpoints": self.traffic.get("cloud_endpoints", [])}] \
                                if self.traffic.get("cloud_endpoints") else \
                                self._match_proc(KW["cloud"] + KW["upload_cmd"])
        idle_exfil = []
        for p in self.c.processes:
            if p["status"] in ("sleeping", "idle"):
                for conn in self.c.connections:
                    if conn["pid"] == p["pid"] and conn["remote"] and conn["status"] == "ESTABLISHED":
                        idle_exfil.append({"type": "idle_network", "pid": p["pid"],
                                           "name": p["name"], "remote": conn["remote"]})
        S["data_idle_exfil"] = idle_exfil
        S["data_stream"]     = self._match_proc(KW["telemetry"])
        if self.traffic.get("outbound_count", 0) > 80:
            S["data_stream"].append({"type": "high_outbound",
                                     "count": self.traffic["outbound_count"]})
        S["data_dns_tunnel"] = self._match_proc(KW["dns_tunnel"])
        S["data_steganography"]= self._match_proc(["steghide", "openstego", "steg_",
                                                     "outguess", "snow_steg"])
        S["data_usb"]        = self._match_proc(["usb_copy", "mass_storage",
                                                   "removable", "udisk"])
        S["data_backup_exfil"]= self._match_files(KW["backup"])

        # ── Model Domain ──
        S["model_context"]   = [{"type": "model_proc", **m}
                                for m in self.model.get("model_procs", [])]
        S["model_prompt"]    = [{"type": "prompt_file", "path": p}
                                for p in self.model.get("prompt_files", [])]
        S["model_finetune"]  = [{"type": "finetune_proc", **m}
                                for m in self.model.get("finetune_procs", [])]
        S["model_embedding"] = [{"type": "embedding_proc", **m}
                                for m in self.model.get("embedding_procs", [])]
        S["model_api_leak"]  = [{"type": "api_endpoint", "endpoint": ep}
                                for ep in self.model.get("api_connections", [])]

        # ── Audit Domain ──
        S["audit_system"]    = self._match_proc(KW["audit"]) + self._match_files(KW["audit"])
        S["audit_log_tamper"]= self._match_proc(["log_delete", "log_truncate",
                                                   "shred", "wipe_log"])
        # Compliance: check if any data protection markers exist
        S["audit_compliance"]= []  # passive — flagged if no audit + data risks
        S["audit_leak_risk"] = []  # computed below

        # Composite leak risk
        active = sum(1 for k, v in S.items()
                     if v and not k.startswith("audit"))
        if active >= 5:
            S["audit_leak_risk"].append({
                "type": "composite",
                "active_categories": active,
                "detail": f"{active} 个安全类目同时触发，数据泄露风险极高",
            })
        if not S["audit_system"] and active >= 2:
            S["audit_compliance"].append({
                "type": "no_audit_with_risks",
                "detail": "无安全审计体系但检测到多项风险信号",
            })

        return S


# ━━━━━━━━━━━━━━━━━━ Risk Engine ━━━━━━━━━━━━━━━━━━

class RiskEngine:
    """加权评分 + 置信度"""

    def evaluate(self, signals: dict) -> list[dict]:
        results = []
        for check_id, domain, label, desc in CHECK_ITEMS:
            evidence = signals.get(check_id, [])
            domain_weight = DOMAINS[domain]["weight"]

            # 审计类特殊逻辑
            if check_id == "audit_system":
                detected = len(evidence) > 0
                raw_score = 0 if detected else 50
                status = "是" if detected else "否（缺失）"
            elif check_id in ("audit_compliance", "audit_leak_risk"):
                detected = len(evidence) > 0
                raw_score = min(100, len(evidence) * 40) if detected else 0
                status = "是（风险）" if detected else "否"
            elif check_id == "audit_log_tamper":
                detected = len(evidence) > 0
                raw_score = min(100, len(evidence) * 50) if detected else 0
                status = "是" if detected else "否"
            else:
                detected = len(evidence) > 0
                raw_score = min(100, len(evidence) * 20 + 10) if detected else 0
                status = "是" if detected else "否"

            weighted_score = min(100, int(raw_score * domain_weight))
            confidence = min(1.0, 0.3 + len(evidence) * 0.15) if detected else 0.95

            results.append({
                "id": check_id,
                "domain": domain,
                "domain_name": DOMAINS[domain]["name"],
                "domain_icon": DOMAINS[domain]["icon"],
                "label": label,
                "description": desc,
                "status": status,
                "detected": detected,
                "risk_score": weighted_score,
                "confidence": round(confidence, 2),
                "evidence_count": len(evidence),
                "evidence": evidence[:8],
            })
        return results


# ━━━━━━━━━━━━━━━━━━ Report Writer ━━━━━━━━━━━━━━━━━━

class ReportWriter:
    """输出 JSON + DOCX 分类报告"""

    def __init__(self, results, output_dir):
        self.results = results
        self.out = Path(output_dir)
        self.out.mkdir(parents=True, exist_ok=True)
        self.ts = SCAN_TS.strftime("%Y-%m-%d %H:%M:%S")

    def _domain_summary(self):
        by_domain = defaultdict(list)
        for r in self.results:
            by_domain[r["domain"]].append(r)
        summary = {}
        for dom, items in by_domain.items():
            risks = [i for i in items if i["detected"]
                     and i["id"] not in ("audit_system",)]
            max_s = max((i["risk_score"] for i in items), default=0)
            summary[dom] = {
                "name": DOMAINS[dom]["name"],
                "icon": DOMAINS[dom]["icon"],
                "total": len(items),
                "risks": len(risks),
                "max_score": max_s,
            }
        return summary

    def write_json(self) -> str:
        ds = self._domain_summary()
        total_risks = sum(d["risks"] for d in ds.values())
        report = {
            "report": "Police Claw v3.0 Enterprise Scan Report",
            "version": VERSION,
            "scan_id": SCAN_ID,
            "timestamp": self.ts,
            "host": platform.node(),
            "os": f"{platform.system()} {platform.release()}",
            "summary": {
                "total_checks": len(self.results),
                "total_risks": total_risks,
                "max_risk_score": max((r["risk_score"] for r in self.results), default=0),
                "domain_summary": ds,
            },
            "checks": self.results,
        }
        p = self.out / "Police_Claw_v3_Report.json"
        p.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
        return str(p)

    def write_docx(self) -> str | None:
        if not HAS_DOCX:
            return None

        doc = DocxDocument()

        # Title
        t = doc.add_heading("Police Claw v3.0", level=0)
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in t.runs:
            r.font.color.rgb = RGBColor(0x0F, 0x0F, 0x23)

        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = sub.add_run("AI Agent / Variant Lobster / Skills 全域安全审计报告")
        r.font.size = Pt(13)
        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        # Meta
        doc.add_paragraph("")
        meta_tbl = doc.add_table(rows=5, cols=2)
        meta_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        meta_data = [
            ("扫描 ID", SCAN_ID), ("时间", self.ts),
            ("主机", platform.node()),
            ("系统", f"{platform.system()} {platform.release()}"),
            ("检查项", f"{len(self.results)} 项 / 7 大安全域"),
        ]
        for i, (k, v) in enumerate(meta_data):
            meta_tbl.cell(i, 0).text = k
            meta_tbl.cell(i, 1).text = v

        # Domain summary
        doc.add_heading("安全域总览", level=1)
        ds = self._domain_summary()
        dtbl = doc.add_table(rows=1, cols=5)
        dtbl.style = "Light Grid Accent 1"
        for i, h in enumerate(["安全域", "检查项数", "风险项数", "最高分", "状态"]):
            c = dtbl.rows[0].cells[i]
            c.text = h
            for p in c.paragraphs:
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(9)
        for dom_id in DOMAINS:
            d = ds.get(dom_id, {})
            row = dtbl.add_row().cells
            row[0].text = f"{d.get('icon','')} {d.get('name','')}"
            row[1].text = str(d.get("total", 0))
            row[2].text = str(d.get("risks", 0))
            row[3].text = str(d.get("max_score", 0))
            status = "安全" if d.get("risks", 0) == 0 else "⚠ 有风险"
            row[4].text = status
            for c in row:
                for p in c.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(9)

        # Per-domain detail
        for dom_id, dom_info in DOMAINS.items():
            doc.add_heading(
                f"{dom_info['icon']} {dom_info['name']}", level=1)
            items = [r for r in self.results if r["domain"] == dom_id]
            tbl = doc.add_table(rows=1, cols=5)
            tbl.style = "Light Grid Accent 1"
            for i, h in enumerate(["检查项", "状态", "风险分", "置信度", "证据数"]):
                c = tbl.rows[0].cells[i]
                c.text = h
                for p in c.paragraphs:
                    for r in p.runs:
                        r.bold = True
                        r.font.size = Pt(9)
            for item in items:
                row = tbl.add_row().cells
                row[0].text = item["label"]
                row[1].text = item["status"]
                row[2].text = str(item["risk_score"])
                row[3].text = f"{item['confidence']:.0%}"
                row[4].text = str(item["evidence_count"])
                if item["risk_score"] >= 50:
                    for p in row[1].paragraphs:
                        for r in p.runs:
                            r.font.color.rgb = RGBColor(0xCC, 0, 0)
                for c in row:
                    for p in c.paragraphs:
                        for r in p.runs:
                            r.font.size = Pt(9)

            # Evidence
            flagged = [it for it in items if it["evidence"]]
            if flagged:
                for it in flagged:
                    doc.add_heading(it["label"], level=2)
                    doc.add_paragraph(it["description"]).runs[0].font.size = Pt(9)
                    for ev in it["evidence"][:5]:
                        para = doc.add_paragraph(style="List Bullet")
                        txt = json.dumps(ev, ensure_ascii=False)
                        if len(txt) > 250:
                            txt = txt[:250] + "..."
                        run = para.add_run(txt)
                        run.font.size = Pt(8)
                        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

        # Risk summary
        doc.add_heading("风险总结", level=1)
        total_risks = sum(1 for r in self.results if r["detected"]
                          and r["id"] not in ("audit_system",))
        mx = max((r["risk_score"] for r in self.results), default=0)
        p = doc.add_paragraph()
        p.add_run(f"共 {len(self.results)} 项检查 / 7 大安全域，")
        p.add_run(f"发现 {total_risks} 项风险，最高风险分 {mx}/100。")
        if total_risks == 0:
            doc.add_paragraph("当前环境未检测到变体龙虾/智能体越权行为。")
        elif total_risks >= 8:
            p2 = doc.add_paragraph()
            r = p2.add_run("⚠ 检测到大量安全风险，强烈建议立即隔离相关进程并执行全面审计。")
            r.bold = True
            r.font.color.rgb = RGBColor(0xCC, 0, 0)

        # Footer
        doc.add_paragraph("")
        ft = doc.add_paragraph()
        ft.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = ft.add_run(f"Police Claw v{VERSION} — {self.ts}")
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        path = self.out / "Police_Claw_v3_Report.docx"
        doc.save(str(path))
        return str(path)


# ━━━━━━━━━━━━━━━━━━ Main ━━━━━━━━━━━━━━━━━━

def main():
    out = sys.argv[1] if len(sys.argv) > 1 else "./police_claw_output"

    print(f"""
╔═══════════════════════════════════════════════════════╗
║         🦞  Police Claw v{VERSION}  Enterprise          ║
║   AI Agent / Variant Lobster / Skills 全域安全审计    ║
║          42 项检查 · 7 大安全域 · 分类报告            ║
╚═══════════════════════════════════════════════════════╝
    """)
    print(f"  Scan ID : {SCAN_ID}")
    print(f"  Host    : {platform.node()} ({platform.system()} {platform.release()})")
    print(f"  Output  : {out}")
    print()

    # Phase 1: Collect
    print("[1/6] 采集系统数据...")
    collector = Collector().collect_all()
    print(f"      ├─ 进程: {len(collector.processes)}")
    print(f"      ├─ 连接: {len(collector.connections)}")
    print(f"      ├─ 文件: {len(collector.open_files)}")
    print(f"      ├─ 环境: {len(collector.env_signals)}")
    print(f"      └─ DNS:  {collector.dns_servers}")

    # Phase 2: Traffic
    print("[2/6] 网络流量分析...")
    traffic = TrafficMonitor(collector).analyze()
    print(f"      ├─ 外传连接: {traffic['outbound_count']}")
    print(f"      ├─ 云端端点: {len(traffic['cloud_endpoints'])}")
    print(f"      └─ 模型API:  {len(traffic['model_api_endpoints'])}")

    # Phase 3: FS
    print("[3/6] 文件系统监控...")
    fs_data = FSMonitor(collector).analyze()
    zones_hit = sum(1 for v in fs_data.values() if v)
    print(f"      └─ 敏感区域命中: {zones_hit}/{len(FSMonitor.SENSITIVE_ZONES)}")

    # Phase 4: Model
    print("[4/6] 模型 API 检测...")
    model_data = ModelMonitor(collector, traffic).analyze()
    print(f"      ├─ AI 进程:   {len(model_data['model_procs'])}")
    print(f"      ├─ Embedding:  {len(model_data['embedding_procs'])}")
    print(f"      └─ Prompt 文件: {len(model_data['prompt_files'])}")

    # Phase 5: Signal + Risk
    print("[5/6] 信号识别 & 风险评估...")
    signals = SignalEngine(collector, traffic, fs_data, model_data).analyze()
    active = sum(1 for v in signals.values() if v)
    print(f"      └─ 活跃信号: {active}/{len(signals)}")

    results = RiskEngine().evaluate(signals)
    risks = sum(1 for r in results if r["detected"]
                and r["id"] not in ("audit_system",))
    print(f"      └─ 风险项: {risks}/{len(results)}")

    # Phase 6: Report
    print("[6/6] 生成报告...")
    writer = ReportWriter(results, out)
    jp = writer.write_json()
    print(f"      ├─ JSON: {jp}")
    dp = writer.write_docx()
    print(f"      └─ DOCX: {dp or 'skipped'}")

    # Summary
    print()
    print("═" * 58)
    current_domain = ""
    for r in results:
        if r["domain"] != current_domain:
            current_domain = r["domain"]
            di = DOMAINS[current_domain]
            print(f"\n  {di['icon']} {di['name']}")
            print(f"  {'─' * 48}")
        is_risk = r["detected"] and r["id"] not in ("audit_system",)
        is_audit_missing = r["id"] == "audit_system" and not r["detected"]
        if is_risk or is_audit_missing:
            icon = "🔴"
        elif r["id"] == "audit_system" and r["detected"]:
            icon = "🟢"
        else:
            icon = "🟢"
        score_str = f"[{r['risk_score']:>3}]" if r["risk_score"] > 0 else "[ 0 ]"
        print(f"    {icon} {r['status']:<8} {score_str}  {r['label']}")

    print(f"\n{'═' * 58}")
    mx = max((r["risk_score"] for r in results), default=0)
    print(f"  合计: {len(results)} 项检查 | {risks} 项风险 | 最高分 {mx}/100")
    print()


if __name__ == "__main__":
    main()
